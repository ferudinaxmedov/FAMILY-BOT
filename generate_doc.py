"""ARIZA hujjatini python-docx bilan yaratadi"""
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

APT_CFG = {
    "23":  {"ishonchnoma": "26.06.2028", "cad": "023"},
    "28":  {"ishonchnoma": "25.06.2028", "cad": "028"},
    "68":  {"ishonchnoma": "25.06.2028", "cad": "068"},
    "80":  {"ishonchnoma": "25.06.2028", "cad": "080"},
    "84":  {"ishonchnoma": "25.06.2028", "cad": "084"},
    "88":  {"ishonchnoma": "25.06.2028", "cad": "088"},
    "701": {"ishonchnoma": "25.06.2028", "cad": "701"},
}
CAD_BASE = "10.05.03.01.05.5583.0001"

def to_disp(iso_date):
    if not iso_date: return "_______"
    try:
        y, m, d = str(iso_date).split("-")
        return f"{d}.{m}.{y}"
    except: return str(iso_date)

def _remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "none")
                tc_borders.append(b)
            tc_pr.append(tc_borders)

def _set_font(run, size=13, bold=False, underline=False):
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"

def _add_header(doc, apt):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(9)
    _remove_table_borders(table)
    lines = [
        "TOSHKENT SHAHRI", "YAKKASAROYTUMANI ИИБ",
        "ФМБ М VA РБ БОШЛИҒИГА", "",
        "TOSHKENT SHAXRI", "ФАРҒОНА ЙО'ЛИ КО'ЧАСИ",
        "25/14, 70-ХОНАДОНДА", "ЯШОВЧИ 18.10.2005 ДА",
        "ТУГ'УЛГАН АХМЕДОВ",
        "ФЕРУDИН САМАД О'ҒЛИДАН",
        "AD0863367 ОЛОТ ТУМАНИ ИИБ",
        "ТОМОНИДАН 29.11.2021 ЙИЛДА",
        "БЕРИЛГАН ТЕЛЕФОН 93 245 06 25",
    ]
    right_cell = table.cell(0, 1)
    first = True
    for line in lines:
        p = right_cell.paragraphs[0] if first else right_cell.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if line:
            run = p.add_run(line)
            _set_font(run, size=11, bold=True)

def generate_ariza_doc(registration):
    apt = str(registration.get("apartment", "28"))
    cfg = APT_CFG.get(apt, {"ishonchnoma": "25.06.2028", "cad": apt.zfill(3)})
    cad = f"{CAD_BASE}.{cfg['cad']}"
    guests = registration.get("guests", [{}])
    if not guests:
        guests = [{"name": "", "nationality": "", "dob": "", "passportId": ""}]

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    _add_header(doc, apt)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(4)
    sp.paragraph_format.space_after = Pt(4)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(14)
    tr = title_p.add_run("ARIZA")
    tr.bold = True; tr.font.size = Pt(16); tr.font.name = "Times New Roman"

    guest_parts = []
    for g in guests:
        dob_disp = to_disp(g.get("dob", ""))
        nationality = g.get("nationality", "")
        name = g.get("name", "")
        guest_parts.append(f"{dob_disp} da tug'ulgan {nationality} fuqarosi {name}")
    guests_str = ",\n".join(guest_parts)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.paragraph_format.space_after = Pt(10)
    r1 = p1.add_run(
        f"Men ariza orqali harakat qiluvchi Ahmedov Ferudin Samad o'g'li "
        f"ushbu arizam bilan shuni malum qilamanki ishonchnommaga asosan Toshkent "
        f"shahri Yakkasaroy tumani Rakatboshi M.F.Y, Bahodir 2/2 uy, "
        f"{apt} xonadonga {guests_str} "
        f"vaqtinchalik ro'yxatga qo'yishingizni so'rayman."
    )
    _set_font(r1, underline=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.first_line_indent = Cm(1.25)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(
        f"Kelgusida ushbu xorijiy fuqaroni vaqtinchalik ro'yxatga "
        f"qo'yilgani yuzasidan hech qanday etirozim va qarshilgim "
        f"bo'lmaydi. Arizaga ishonchnoma kseronusxasini ilova qildim. "
        f"Ishonchnoma {cfg['ishonchnoma']} muddatigacha berilgan"
    )
    _set_font(r2, underline=True)

    p_cad = doc.add_paragraph()
    p_cad.paragraph_format.space_after = Pt(4)
    r_cad = p_cad.add_run(f"Kadastr raqami: {cad}")
    _set_font(r_cad, underline=True)

    for _ in range(4): doc.add_paragraph()

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = False
    sig_table.columns[0].width = Cm(9)
    sig_table.columns[1].width = Cm(8)
    _remove_table_borders(sig_table)
    lc = sig_table.cell(0, 0); rc = sig_table.cell(0, 1)
    lc.paragraphs[0].clear()
    lr = lc.paragraphs[0].add_run("Ahmedov Ferudin Samad o'g'li")
    _set_font(lr, underline=True)
    rc.paragraphs[0].clear()
    rr = rc.paragraphs[0].add_run("«    »                    2026")
    _set_font(rr)
    rc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("IMZO")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
